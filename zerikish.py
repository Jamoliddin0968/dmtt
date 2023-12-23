from docx2pdf import convert
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# Yangi Word hujjat yaratish
doc = Document()

# ------------------------------------------------------
paragraph = doc.add_paragraph()
text = "06.09.2023 sanadagi 128-sonli shartnomaga\n 01.09-27.09 2023 sanalardagi 03/09-sonli\n"
font = paragraph.add_run(text).font
font.name = 'Times New Roman'
font.size = Pt(11)
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# ------------------------------------------
paragraph_text = "HISOBVARAQ-FAKTURA"
paragraph = doc.add_paragraph()
run = paragraph.add_run(paragraph_text)

font = run.font
font.name = 'Times New Roman'
font.size = Pt(19)
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# -----------------------------------------------------------


data = [
    [
        {"Key": "Yetkazib beruvchi:", "Value": "Muhammadqodir"},
        {"Key": "Xaridor:", "Value": "df"},
    ],
    [
        {"Key": "Manzil:", "Value": "Toshloq"},
        {"Key": "Manzil:", "Value": "None"},
    ],
    [
        {"Key": "Tel:", "Value": "+9983698008"},
        {"Key": "Tel:", "Value": "none"},
    ],
    [
        {"Key": "STIR:", "Value": "7895152"},
        {"Key": "STIR:", "Value": "None"}
    ]
]

# Add a table with 4 columns and 4 rows without lines and borders
columns = 4
rows = 4
table = doc.add_table(rows=rows, cols=columns)
table.style = 'Table Grid'


for i in range(rows):
    for j in range(columns):
        cell = table.cell(i, j)
        key = data[i][j//2]["Key"]
        value = data[i][j//2]["Value"]
        # Set styles for keys (column 1) and values (column 2)
        if j % 2 == 0:  # If the cell is in the left column
            key_run = cell.paragraphs[0].add_run(key)
            key_run.font.name = 'Times New Roman'
            key_run.font.size = Pt(11)
            key_run.font.bold = True
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:  # If the cell is in the right column
            value_run = cell.paragraphs[0].add_run(
                value if value else "")  # Add value or empty string if None
            value_run.font.name = 'Times New Roman'
            value_run.font.size = Pt(11)


# ---------------------------------------------------------


doc.save('mavzu.docx')
convert("mavzu.docx", "output.pdf")
